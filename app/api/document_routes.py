"""
document_routes.py — File-upload endpoint for scanning PDF and DOC/DOCX documents.

Accepts up to 5 files via multipart/form-data, extracts their text content,
and runs the same security scan pipeline as /secure-retrieve.
"""

import io
import uuid
import time
import hashlib
import logging
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile, status
from sqlalchemy.orm import Session

from app.auth.api_key import verify_api_key
from app.models.schemas import SecureRetrieveResponse, ThreatDetail
from app.core.security_engine import scan_query, scan_chunks
from app.core.scoring import compute_score
from app.ml.classifier import predict as ml_predict
from app.db.database import (
    get_db,
    save_request,
    save_threats,
    save_score_audit,
    get_demo_usage,
    increment_demo_usage,
    DEMO_LIMIT,
)

logger = logging.getLogger("securerag.documents")

document_router = APIRouter()

MAX_FILES = 5
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


# ── Text extraction helpers ──────────────────────────────────────────────────

def _extract_text_pdf(data: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text.strip())
        doc.close()
        return "\n\n".join(pages)
    except Exception as e:
        logger.warning("PDF extraction failed: %s", e)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Could not extract text from PDF: {e}",
        )


def _extract_text_docx(data: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in paragraphs:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Could not extract text from DOCX: {e}",
        )


def _extract_text_doc(data: bytes) -> str:
    """Extract text from legacy .doc file via LibreOffice conversion fallback."""
    try:
        import subprocess
        import tempfile
        import os
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = os.path.join(tmpdir, "input.doc")
            with open(doc_path, "wb") as f:
                f.write(data)
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, doc_path],
                capture_output=True, timeout=30,
            )
            txt_path = os.path.join(tmpdir, "input.txt")
            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read().strip()
        raise RuntimeError("LibreOffice conversion produced no output")
    except Exception as e:
        logger.warning("DOC extraction failed: %s", e)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Could not extract text from legacy .doc file: {e}. Try converting to .docx first.",
        )


async def _read_and_extract(file: UploadFile) -> str:
    """Read uploaded file, validate it, and extract text."""
    filename = file.filename or "unknown"
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"File '{filename}': unsupported type '{ext}'. Allowed: .pdf, .doc, .docx",
        )

    data = await file.read()

    if len(data) == 0:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"File '{filename}' is empty.",
        )

    if len(data) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File '{filename}' exceeds {MAX_FILE_SIZE_MB} MB limit.",
        )

    if ext == ".pdf":
        return _extract_text_pdf(data)
    elif ext == ".docx":
        return _extract_text_docx(data)
    elif ext == ".doc":
        return _extract_text_doc(data)

    raise HTTPException(status_code=400, detail=f"Unhandled extension: {ext}")


# ── Shared scan logic (mirrors routes.py _run_scan but accepts pre-extracted chunks) ──

def _run_document_scan(
    user_query: str,
    file_chunks: List[str],
    db: Session,
    api_key_hint: Optional[str],
):
    request_id = str(uuid.uuid4())
    t_start = time.perf_counter()

    query_threats = scan_query(user_query)
    chunk_threats, unsafe_chunk_indices = scan_chunks(file_chunks)
    all_threats = query_threats + chunk_threats
    ml_result = ml_predict(user_query)
    risk_score, component_scores, blocked, reasons = compute_score(
        all_threats, ml_result=ml_result
    )

    safe_chunks: List[str] = [
        chunk for i, chunk in enumerate(file_chunks)
        if i not in unsafe_chunk_indices
    ]
    if blocked and (query_threats or ml_result.get("malicious")):
        safe_chunks = []

    processing_time_ms = (time.perf_counter() - t_start) * 1000
    query_hash = hashlib.sha256(user_query.encode()).hexdigest()

    save_request(
        db=db, request_id=request_id, user_id=None,
        query_hash=query_hash, risk_score=risk_score, blocked=blocked,
        chunks_received=len(file_chunks), chunks_passed=len(safe_chunks),
        threats_detected=len(all_threats), processing_time_ms=processing_time_ms,
        api_key_hint=api_key_hint,
    )
    if all_threats:
        save_threats(db=db, request_id=request_id, threats=all_threats)
    component_scores["ml_result"] = ml_result
    save_score_audit(db=db, request_id=request_id,
                     component_scores=component_scores, final_score=risk_score)

    threats = [
        ThreatDetail(
            threat_type=t["threat_type"], description=t["description"],
            severity=t["severity"], matched_pattern=t.get("matched_pattern"),
            source=t["source"],
        )
        for t in all_threats
    ]

    return dict(
        request_id=request_id, safe_chunks=safe_chunks, risk_score=risk_score,
        blocked=blocked, reasons=reasons, threats=threats,
        chunks_filtered=len(file_chunks) - len(safe_chunks),
        processing_time_ms=round(processing_time_ms, 2),
        timestamp=datetime.now(timezone.utc),
    )


# ── Paid endpoint ─────────────────────────────────────────────────────────────

@document_router.post("/scan-files", response_model=SecureRetrieveResponse)
async def scan_files(
    user_query: str = Form(..., min_length=1, max_length=10000,
                           description="User query or message to scan"),
    files: List[UploadFile] = File(..., description="1–5 PDF or DOC/DOCX files to scan"),
    db: Session = Depends(get_db),
    api_key: str = Depends(verify_api_key),
):
    """
    Scan a user query **plus** up to 5 uploaded PDF/DOC/DOCX files for prompt injection
    and indirect document poisoning threats.

    - Files are parsed server-side; extracted text is treated as retrieved chunks.
    - Same detection pipeline as /secure-retrieve (regex + ML).
    - Requires a valid API key.
    """
    if len(files) > MAX_FILES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Too many files. Maximum is {MAX_FILES}.",
        )

    file_chunks: List[str] = []
    for f in files:
        text = await _read_and_extract(f)
        if text:
            file_chunks.append(text)
        else:
            file_chunks.append(f"(file '{f.filename}' contained no extractable text)")

    if not file_chunks:
        file_chunks = ["(no documents provided)"]

    result = _run_document_scan(
        user_query=user_query,
        file_chunks=file_chunks,
        db=db,
        api_key_hint=(api_key[:8] + "…") if api_key else None,
    )
    return SecureRetrieveResponse(**result)


# ── Demo endpoint ─────────────────────────────────────────────────────────────

def _ip_hash(request: Request) -> str:
    ip = request.headers.get("x-forwarded-for", request.client.host or "unknown")
    ip = ip.split(",")[0].strip()
    return hashlib.sha256(ip.encode()).hexdigest()


@document_router.post("/demo/scan-files", response_model=SecureRetrieveResponse)
async def demo_scan_files(
    request: Request,
    user_query: str = Form(..., min_length=1, max_length=10000),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    """
    Demo version of /scan-files — no API key required.
    Subject to the same IP-based demo limit as /demo/scan.
    """
    ip_hash = _ip_hash(request)
    used = get_demo_usage(db, ip_hash)

    if used >= DEMO_LIMIT:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail={
                "error": "demo_limit_exhausted",
                "message": f"Demo limit of {DEMO_LIMIT} scans reached. Subscribe to continue.",
                "scans_used": used,
                "scans_limit": DEMO_LIMIT,
            },
        )

    if len(files) > MAX_FILES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail=f"Too many files. Maximum is {MAX_FILES}.",
        )

    file_chunks: List[str] = []
    for f in files:
        text = await _read_and_extract(f)
        if text:
            file_chunks.append(text)
        else:
            file_chunks.append(f"(file '{f.filename}' contained no extractable text)")

    if not file_chunks:
        file_chunks = ["(no documents provided)"]

    result = _run_document_scan(
        user_query=user_query,
        file_chunks=file_chunks,
        db=db,
        api_key_hint="demo",
    )
    new_count = increment_demo_usage(db, ip_hash)

    return SecureRetrieveResponse(
        **result,
        demo_scans_used=new_count,
        demo_scans_remaining=max(0, DEMO_LIMIT - new_count),
    )
